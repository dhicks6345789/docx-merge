#!/usr/bin/python

# DOCX Merge - a utility to do document merges with DOCX files.

# Standard libraries.
import os
import re
import sys
import math
import shutil
import zipfile
import datetime

# The python-docx library, for manipulating DOCX files.
# Importantly, when installing with pip, that not the "docx" library, that an earlier version - do "pip install python-docx".
import docx

# We use Pandas to import Excel / CSV files for configuration details.
import pandas

# The pytz library, for dealing sensibly with timezones.
import pytz
theTimezone = pytz.timezone("Europe/London")

# Possible states for the iCal parser.
ICALSTART = 0
ICALINVEVENT = 1

# Used for placeholder names in calendars.
DAYNAMES = ["MO","TU","WE","TH","FR","SA","SU"]

# DOCX files are ZIP files - we need a folder to unzip the contenst into if we want to modify a contained file.
TEMPLATETEMP = "templateTemp/"

ONEDAY = datetime.timedelta(days=1)

# A place to put calendar data read from an iCal file.
calendar = {}

# Make sure the given Year exists in the calendar.
def addCalendarYear(theYear):
	if not theYear in calendar.keys():
		calendar[theYear] = {}

# Make sure the given Month exists in the calendar.
def addCalendarMonth(theYear, theMonth):
	addCalendarYear(theYear)
	if not theMonth in calendar[theYear].keys():
		calendar[theYear][theMonth] = {}

# Make sure the given Day exists in the calendar.		
def addCalendarDay(theYear, theMonth, theDay):
	addCalendarMonth(theYear, theMonth)
	if not theDay in calendar[theYear][theMonth].keys():
		calendar[theYear][theMonth][theDay] = []

# Add the given Item to the calendar, at the given day / month / year.
def addCalendarItem(theYear, theMonth, theDay, theItem):
	addCalendarDay(theYear, theMonth, theDay)
	calendar[theYear][theMonth][theDay].append(theItem)

# Strip unwanted characters from strings.
def normaliseString(theString):
	result = ""
	for resultItem in theString.replace("\\n","\n").replace("\\,",",").replace("·","").split("\n"):
		result = result + resultItem.strip() + "\n"
	return(result.strip())

def unZeroPad(theString):
	if theString[0] == "0":
		return(theString[1:])
	return(theString)

def time24To12Hour(theString):
	result = theString
	matchResult = re.match("(\d\d):(\d\d): ", theString)
	if not matchResult == None:
		hour = int(matchResult.group(1))
		minuteString = ""
		if not matchResult.group(2) == "00":
			minuteString = ":" + str(int(matchResult.group(2)))
		if hour == 12 and matchResult.group(0) == "00":
			result = "Midday: "
		elif hour > 12:
			result = str(hour-12) + minuteString + "pm: "
		else:
			result = str(hour) + minuteString + "am: "
		result = result + theString[7:]
	return(result)

# A basic iCal parser.
def parseICalFile(theFilename):
	iCalState = ICALSTART
	iCalData = {}
	
	# Read the iCal file in as a bunch of text entries. We don't just use readlines() as some entries can be split over multiple lines, so we have to detect
	# those and stick them back together as we go along.
	iCalLines = []
	iCalHandle = open(theFilename, encoding="utf-8")
	for iCalLine in iCalHandle:
		if iCalLine.startswith(" "):
			iCalLines[-1] = iCalLines[-1] + iCalLine[1:].rstrip()
		else:
			iCalLines.append(iCalLine.rstrip())
	iCalHandle.close()
	
	# Now, parse the lines read above into calendar data.
	iCalBlock = ""
	for iCalLine in iCalLines:
		if iCalState == ICALSTART and iCalLine.startswith("BEGIN:VEVENT"):
			iCalState = ICALINVEVENT
			iCalData = {}
			iCalBlock = ""
		if iCalState == ICALINVEVENT:
			iCalBlock = iCalBlock + iCalLine + "\n"
		if iCalState == ICALINVEVENT and iCalLine.startswith("DTSTART:"):
			startDateTime = theTimezone.localize(datetime.datetime.strptime(iCalLine.split(":",1)[1], "%Y%m%dT%H%M%SZ"))
			iCalData["StartDate"] = startDateTime + startDateTime.utcoffset()
			iCalData["StartTime"] = startDateTime + startDateTime.utcoffset()
		if iCalState == ICALINVEVENT and iCalLine.startswith("DTSTART;VALUE=DATE:"):
			iCalData["StartDate"] = datetime.datetime.strptime(iCalLine.split(":",1)[1], "%Y%m%d")
		if iCalState == ICALINVEVENT and iCalLine.startswith("DTEND:"):
			endDateTime = theTimezone.localize(datetime.datetime.strptime(iCalLine.split(":",1)[1], "%Y%m%dT%H%M%SZ"))
			iCalData["EndDate"] = endDateTime + startDateTime.utcoffset()
			iCalData["EndTime"] = endDateTime + startDateTime.utcoffset()
		if iCalState == ICALINVEVENT and iCalLine.startswith("DTEND;VALUE=DATE:"):
			iCalData["EndDate"] = datetime.datetime.strptime(iCalLine.split(":",1)[1], "%Y%m%d")
		if iCalState == ICALINVEVENT and iCalLine.startswith("SUMMARY:"):
			summary = iCalLine.split(":",1)[1].strip()
			if not description == "":
				iCalData["Summary"] = summary
		if iCalState == ICALINVEVENT and iCalLine.startswith("DESCRIPTION:"):
			description = iCalLine.split(":",1)[1].strip()
			if not description == "":
				iCalData["Description"] = description
		if iCalState == ICALINVEVENT and iCalLine.startswith("LOCATION:"):
			location = iCalLine.split(":",1)[1].strip()
			if not location == "":
				iCalData["Location"] = location
		if iCalState == ICALINVEVENT and iCalLine.startswith("END:VEVENT"):
			iCalState = ICALSTART
			if "StartDate" in iCalData.keys():
				# Do we not have an EndDate set? Assume the event lasts one day.
				if not "EndDate" in iCalData.keys():
					iCalData["EndDate"] = iCalData["StartDate"]
				# Does the event not have a start / end time set but seems to last from one day until the next? Simply the event into lasting one day.
				if iCalData["EndDate"] == (iCalData["StartDate"] + ONEDAY):
					if not "StartTime" in iCalData.keys() and not "EndTime" in iCalData.keys():
						iCalData["EndDate"] = iCalData["StartDate"]
				timeString = ""
				if "StartTime" in iCalData.keys():
					timeString = iCalData["StartTime"].strftime("%H:%M: ")
				eventLength = iCalData["EndDate"] - iCalData["StartDate"]
				currentDate = iCalData["StartDate"]
				for eventDay in range(0, eventLength.days+1):
					itemText = ""
					if "Summary" in iCalData.keys():
						if not iCalData["Summary"] == iCalData["Description"]:
							itemText = itemText + iCalData["Summary"] + ", "
					itemText = itemText + iCalData["Description"]
					if "Location" in iCalData.keys():
						itemText = itemText + ", " + iCalData["Location"]
					addCalendarItem(currentDate.year, currentDate.month, currentDate.day, timeString + normaliseString(itemText))
					currentDate = currentDate + ONEDAY
			else:
				print("Unhandled event:\n" + iCalBlock.strip())

# Extract the given DOCX file to a given temporary folder.
# Reads and returns the contents of the "document.xml" contained in the DOCX file.
def extractDocx(theFilename, destinationPath):
	templateDocx = zipfile.ZipFile(theFilename, "r")
	templateDocx.extractall(destinationPath)
	templateDocx.close()
	textHandle = open(destinationPath + "word/document.xml")
	docxText = str(textHandle.read())
	textHandle.close()
	return(docxText)

# Turns the contents of the given folder into a DOCX file. Deletes the source folder when done.
def compressDocx(sourcePath, theFilename):
	theDocx = zipfile.ZipFile(theFilename, "w")
	for root, dirs, files in os.walk(sourcePath):
		for file in files:
			theDocx.write(os.path.join(root, file), os.path.join(root, file)[len(sourcePath):])
	theDocx.close()
	shutil.rmtree(sourcePath)

# Writes a file to the given path.
def putFile(thePath, theData):
	textHandle = open(thePath, "w")
	textHandle.write(theData)
	textHandle.close()
	
def checkForRequiredArgs(theActualArgs, theRequiredArgs):
	for requiredArg in theRequiredArgs:
		if not requiredArg in theActualArgs.keys():
			print("ERROR: argument missing, " + requiredArg)
			sys.exit(1)
			
def cellToStr(theInput):
	if isinstance(theInput, str):
		return(theInput)
	if isinstance(theInput, float) and math.isnan(theInput):
		return("")
	return(str(theInput))

def calendarItemSortOrder(theItem):
	for pl in range(0, len(theItem)):
		for numeral in ["0", "1", "2", "3", "4", "5", "6", "7", "8", "9"]:
			if theItem[pl] == numeral:
				theItem[pl] = char(ord(numeral) + 128)
	return theItem
											      
# Check arguments, print a usage message if needed.
if len(sys.argv) == 1:
	print("DOCX Merge - merges data into DOCX templates. Usage:")
	print("merge.py --week-to-view startDate noOfWeeks data.ics template.docx output.docx")
	sys.exit(0)
	
# Parse command-line arguments.
args = {}
currentArgName = None
for argItem in sys.argv[1:]:
	if argItem.startswith("--"):
		currentArgName = argItem[2:]
	elif not currentArgName == None:
		args[currentArgName] = argItem
		currentArgName = None
	else:
		print("ERROR: unknown argument, " + argItem)
		sys.exit(1)

if "config" in args.keys():
	if args["config"].endswith(".csv"):
		argsData = pandas.read_csv(args["config"], header=0)
	else:
		argsData = pandas.read_excel(args["config"], header=0)
	for argsDataIndex, argsDataValues in argsData.iterrows():
		args[argsDataValues[0]] = cellToStr(argsDataValues[1])

# The user wants a week-to-view calendar.
if "mergeType" in args.keys() and args["mergeType"] == "week-to-view":
	checkForRequiredArgs(args, ["startDate","noOfWeeks","calendar","template","output"])
	
	# Check the start date is a Monday.
	startDate = theTimezone.localize(datetime.datetime.strptime(args["startDate"], "%Y%m%d"))
	if not startDate.weekday() == 0:
		print("ERROR: Start date is not a Monday.")
		sys.exit(1)
	# Figure out the number of weeks to produce a calendar for.
	noOfWeeks = int(args["noOfWeeks"])
	# Read the calendar data.
	parseICalFile(args["calendar"])
	
	# The python-docx library doesn't have a function to duplicate pages, so we do that part ourselves by duplicating the main body of XML from
	# the "document.xml" contained in the DOCX file.
	docxText = extractDocx(args["template"], TEMPLATETEMP)
	bodyStart = docxText.find("<w:body>")+8
	bodyEnd = docxText.find("</w:body>")
	newDocxText = docxText[:bodyStart]
	# Copy the main body text once for each week the user wants to show. Might be multiple pages.
	for week in range(0, noOfWeeks):
		weekToViewText = docxText[bodyStart:bodyEnd]
		for weekDay in range(0, 7):
			today = startDate + datetime.timedelta(days=(week*7)+weekDay)
			# Find the "title" string for the day.
			weekToViewText = weekToViewText.replace("{{" + DAYNAMES[weekDay] + "TI}}", today.strftime("%A, ") + unZeroPad(today.strftime("%d")) + today.strftime(" %B"))
			# Find the "content" string for the day.
			weekToViewText = weekToViewText.replace("{{" + DAYNAMES[weekDay] + "CO}}", "{{" + DAYNAMES[weekDay] + "-WEEK" + str(week) + "}}")
		newDocxText = newDocxText + weekToViewText
	# Re-write the content back to the output location.
	newDocxText = newDocxText + docxText[bodyEnd:]
	putFile(TEMPLATETEMP + "word/document.xml", newDocxText)
	compressDocx(TEMPLATETEMP, args["output"])
		
	# Now, read the output file again with the python-docx library.
	templateDocx = docx.Document(args["output"])
	for week in range(0, noOfWeeks):
		for weekDay in range(0, 7):
			dayContents = ""
			today = startDate + datetime.timedelta(days=(week*7)+weekDay)
			if today.year in calendar.keys():
				if today.month in calendar[today.year].keys():
					if today.day in calendar[today.year][today.month].keys():
						for dayItem in sorted(calendar[today.year][today.month][today.day], key=calendarItemSortOrder):
							dayContents = dayContents + time24To12Hour(dayItem.replace("\n",", ")) + "\n"
			dayContents = dayContents.strip()
			dayString = "{{" + DAYNAMES[weekDay] + "-WEEK" + str(week) + "}}"
			for paragraph in templateDocx.paragraphs:
				for run in paragraph.runs:
					if dayString in run.text:
						run.text = dayContents
			for table in templateDocx.tables:
				for row in table.rows:
					for cell in row.cells:
						for paragraph in cell.paragraphs:
							for run in paragraph.runs:
								if dayString in run.text:
									run.text = dayContents
	# Write out the final version of the DOCX file.
	templateDocx.save(args["output"])
